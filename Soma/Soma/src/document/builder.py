"""
BPDDocumentBuilder — turns a list of Block objects into a polished
.docx file that matches the corporate BPD template.

The builder is the single place where rendering decisions live. Block
classes stay as plain data; the UI only ever hands ordered Blocks to
``build()``.
"""

from __future__ import annotations

import io
from typing import Iterable, List, Optional

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor

from .blocks import (
    Block,
    BlockType,
    HeadingBlock,
    ImageBlock,
    NormalHeadingBlock,
    ParagraphBlock,
    SubheadingBlock,
    TableBlock,
)
from .header_footer import DocumentBranding, add_header_and_footer, configure_page_layout
from .styles import (
    BODY_STYLE,
    HEADING_STYLE,
    NORMAL_HEADING_STYLE,
    SUBHEADING_STYLE,
    TABLE_STYLE,
    register_styles,
)
from .toc import add_table_of_contents, enable_update_fields_on_open


# Corporate grey for table header rows.
HEADER_FILL_HEX = "BFBFBF"
BORDER_HEX = "000000"


class BPDDocumentBuilder:
    """Stateful builder; one instance per generated document."""

    def __init__(self) -> None:
        self.doc: DocxDocument = Document()
        # Counters for auto-numbered headings (e.g. "1", "1.1").
        self._heading_counter = 0
        self._subheading_counter = 0

    # ------------------------------------------------------------------ public

    def build(self, blocks: Iterable[Block], branding: DocumentBranding | None = None) -> bytes:
        """Render the supplied blocks and return the .docx bytes."""
        configure_page_layout(self.doc)
        register_styles(self.doc)
        add_header_and_footer(self.doc, branding)
        add_table_of_contents(self.doc)
        enable_update_fields_on_open(self.doc)

        for block in blocks:
            self._render_block(block)

        buffer = io.BytesIO()
        self.doc.save(buffer)
        return buffer.getvalue()

    # ----------------------------------------------------------------- dispatch

    def _render_block(self, block: Block) -> None:
        if isinstance(block, HeadingBlock):
            self._add_heading(block.text)
        elif isinstance(block, SubheadingBlock):
            self._add_subheading(block.text)
        elif isinstance(block, NormalHeadingBlock):
            self._add_normal_heading(block.text)
        elif isinstance(block, ParagraphBlock):
            self._add_paragraph(block.text)
        elif isinstance(block, ImageBlock):
            self._add_image(block)
        elif isinstance(block, TableBlock):
            self._add_table(block)
        else:
            raise ValueError(f"Unsupported block type: {block.block_type}")

    # ----------------------------------------------------------------- headings

    def _add_heading(self, text: str) -> None:
        self._heading_counter += 1
        self._subheading_counter = 0
        numbered = f"{self._heading_counter} {text.upper()}"
        self.doc.add_paragraph(numbered, style=HEADING_STYLE)

    def _add_subheading(self, text: str) -> None:
        # Subheadings under a heading; if no heading yet, treat as 0.x.
        head = self._heading_counter if self._heading_counter > 0 else 0
        self._subheading_counter += 1
        numbered = f"{head}.{self._subheading_counter} {text.upper()}"
        self.doc.add_paragraph(numbered, style=SUBHEADING_STYLE)

    def _add_normal_heading(self, text: str) -> None:
        # No uppercasing, no TOC, no empty line after.
        self.doc.add_paragraph(text, style=NORMAL_HEADING_STYLE)

    # ---------------------------------------------------------------- paragraph

    def _add_paragraph(self, text: str) -> None:
        # Honour user-entered line breaks by splitting on newlines so each
        # logical line becomes its own paragraph in the same body style.
        for line in (text or "").splitlines() or [""]:
            self.doc.add_paragraph(line, style=BODY_STYLE)

    # -------------------------------------------------------------------- image

    def _add_image(self, block: ImageBlock) -> None:
        if not block.image_bytes:
            return

        # Spacing before the image.
        self.doc.add_paragraph()

        section = self.doc.sections[0]
        content_width_emu = section.page_width - section.left_margin - section.right_margin
        ratio = max(0.1, min(block.width_ratio or 0.7, 1.0))
        target_width = Emu(int(content_width_emu * ratio))

        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        try:
            run.add_picture(io.BytesIO(block.image_bytes), width=target_width)
        except Exception:
            # Bad/unsupported image — drop a placeholder rather than crash
            # the whole document.
            run.add_text(f"[Unable to embed image: {block.filename}]")
            return

        # Black border around the image — applied to the paragraph since
        # python-docx doesn't expose picture borders directly. The border
        # hugs the centered image because the paragraph is the same width.
        self._apply_paragraph_border(para)

        # Spacing after the image.
        self.doc.add_paragraph()

    @staticmethod
    def _apply_paragraph_border(paragraph) -> None:
        pPr = paragraph._p.get_or_add_pPr()
        # Remove any prior border (idempotency).
        for old in pPr.findall(qn("w:pBdr")):
            pPr.remove(old)

        pBdr = OxmlElement("w:pBdr")
        for edge in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "8")  # 1pt (1/8 pt units)
            border.set(qn("w:space"), "4")
            border.set(qn("w:color"), BORDER_HEX)
            pBdr.append(border)
        pPr.append(pBdr)

    # -------------------------------------------------------------------- table

    def _add_table(self, block: TableBlock) -> None:
        block.ensure_shape()
        if block.rows < 1 or block.cols < 1:
            return

        table = self.doc.add_table(rows=block.rows, cols=block.cols)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.autofit = True

        for r in range(block.rows):
            for c in range(block.cols):
                cell = table.cell(r, c)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # Replace the default empty paragraph so we control the style.
                cell.text = ""
                para = cell.paragraphs[0]
                para.style = self.doc.styles[TABLE_STYLE]
                run = para.add_run(block.data[r][c] or "")
                if r == 0:
                    run.bold = True
                    self._shade_cell(cell, HEADER_FILL_HEX)
                self._set_cell_borders(cell)

        # Trailing spacing after the table.
        self.doc.add_paragraph()

    @staticmethod
    def _shade_cell(cell, hex_fill: str) -> None:
        tcPr = cell._tc.get_or_add_tcPr()
        # Remove existing shading so re-runs are idempotent.
        for old in tcPr.findall(qn("w:shd")):
            tcPr.remove(old)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_fill)
        tcPr.append(shd)

    @staticmethod
    def _set_cell_borders(cell) -> None:
        tcPr = cell._tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:tcBorders")):
            tcPr.remove(old)
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")  # 0.5pt
            border.set(qn("w:color"), BORDER_HEX)
            tcBorders.append(border)
        tcPr.append(tcBorders)

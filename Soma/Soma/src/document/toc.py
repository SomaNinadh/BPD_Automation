"""
Table of Contents helpers.

We emit a Word TOC field (rather than a static list of links) so the
TOC stays in sync with the document. The field is configured to pull
outline levels 1-2, which matches our HeadingStyle/SubheadingStyle.
NormalHeadingStyle has no outline level, so it stays out automatically.
"""

from __future__ import annotations

from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


TOC_TITLE = "TABLE OF CONTENTS"


def add_table_of_contents(doc: DocxDocument) -> None:
    """Insert a TOC field and a page break right after it."""
    # TOC title (not part of TOC itself — uses a plain run).
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(TOC_TITLE)
    title_run.bold = True
    title_run.font.size = Pt(16)

    # The TOC field itself.
    p = doc.add_paragraph()
    run = p.add_run()
    run.bold = True

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")  # Force update when the doc opens.

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    # \o "1-2"  → include outline levels 1 and 2
    # \h        → hyperlinked entries
    # \z        → hide tab leader/page numbers in web view
    # \u        → use applied paragraph outline level
    instr.text = r'TOC \o "1-2" \h \z \u'

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose 'Update Field' to refresh the Table of Contents."

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(sep)
    run._r.append(placeholder)
    run._r.append(end)

    # Page break so the body content starts on a fresh page.
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def enable_update_fields_on_open(doc: DocxDocument) -> None:
    """Tell Word to prompt the user (or auto-update) fields on open.

    Without this, the TOC shows our placeholder text until the user
    manually refreshes the field.
    """
    settings = doc.settings.element
    # Avoid duplicate insertion if called twice.
    existing = settings.find(qn("w:updateFields"))
    if existing is not None:
        existing.set(qn("w:val"), "true")
        return
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)

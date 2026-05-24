"""
Header and footer setup.

Builds the branded header/footer for each document section, including
fixed left/right logos, document title, footer metadata, and an
auto-updating page number field.
"""

from __future__ import annotations

import os
from dataclasses import dataclass

from docx.document import Document as DocxDocument
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from pathlib import Path


FONT_NAME = "Calibri"
DEFAULT_PROJECT_NAME = "S4 HANA SANMAR"
# Replace these with the real image paths on the machine/server.
BASE_DIR = Path(__file__).resolve().parents[2]
CLIENT_LOGO_PATH = str(BASE_DIR / "assets" / "logos" / "client-logo.jpg")
OUR_LOGO_PATH = str(BASE_DIR / "assets" / "logos" / "Deloitte-logo.jpg")


@dataclass
class DocumentBranding:
    header_title: str = ""
    module_name: str = ""
    project_name: str = DEFAULT_PROJECT_NAME
    client_logo_path: str = CLIENT_LOGO_PATH
    our_logo_path: str = OUR_LOGO_PATH


def _add_page_number_field(paragraph) -> None:
    """Append a {PAGE} field to the given paragraph."""
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(sep)
    run._r.append(end)


def _clear_story(story) -> None:
    for table in list(story.tables):
        tbl = table._element
        tbl.getparent().remove(tbl)
    for paragraph in list(story.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)


def _set_run_font(run, *, size: int, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold


def _set_cell_width(cell, inches: float) -> None:
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(inches * 1440)))


def _try_add_picture(run, image_source, *, width_inches: float) -> bool:
    if not image_source:
        return False
    try:
        run.add_picture(image_source, width=Inches(width_inches))
        return True
    except Exception:
        return False


def configure_page_layout(doc: DocxDocument) -> None:
    """Set section margins so header/footer have 1.25in of room."""
    for section in doc.sections:
        section.top_margin = Inches(1.25)
        section.bottom_margin = Inches(1.25)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)


def add_header_and_footer(doc: DocxDocument, branding: DocumentBranding | None = None) -> None:
    """Populate every section's header and footer."""
    branding = branding or DocumentBranding()
    header_title = (branding.header_title or "").strip().upper()
    module_name = (branding.module_name or "").strip()
    project_name = (branding.project_name or DEFAULT_PROJECT_NAME).strip()
    client_logo_path = branding.client_logo_path if os.path.exists(branding.client_logo_path) else ""
    our_logo_path = branding.our_logo_path if os.path.exists(branding.our_logo_path) else ""

    for section in doc.sections:
        content_width = section.page_width - section.left_margin - section.right_margin

        header = section.header
        _clear_story(header)
        header_table = header.add_table(rows=1, cols=3, width=content_width)
        header_table.autofit = False

        left_cell, center_cell, right_cell = header_table.rows[0].cells
        for cell in (left_cell, center_cell, right_cell):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        _set_cell_width(left_cell, 1.15)
        _set_cell_width(center_cell, 4.2)
        _set_cell_width(right_cell, 1.15)

        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if not _try_add_picture(left_para.add_run(), client_logo_path, width_inches=0.7):
            fallback = left_para.add_run("CLIENT LOGO")
            _set_run_font(fallback, size=10, bold=True)

        center_para = center_cell.paragraphs[0]
        center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = center_para.add_run(header_title)
        _set_run_font(title_run, size=20, bold=True)

        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if not _try_add_picture(right_para.add_run(), our_logo_path, width_inches=1.1):
            fallback = right_para.add_run("OUR LOGO")
            _set_run_font(fallback, size=10, bold=True)

        footer = section.footer
        _clear_story(footer)
        footer_table = footer.add_table(rows=1, cols=3, width=content_width)
        footer_table.autofit = False

        footer_left, footer_center, footer_right = footer_table.rows[0].cells
        for cell in (footer_left, footer_center, footer_right):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        _set_cell_width(footer_left, 2.4)
        _set_cell_width(footer_center, 0.8)
        _set_cell_width(footer_right, 3.3)

        left_para = footer_left.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_run = left_para.add_run(f"BPD {module_name}".strip())
        _set_run_font(left_run, size=10, bold=True)

        center_para = footer_center.paragraphs[0]
        center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_number_field(center_para)
        if center_para.runs:
            _set_run_font(center_para.runs[-1], size=10, bold=True)

        right_para = footer_right.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_run = right_para.add_run(f"Project: {project_name}")
        _set_run_font(right_run, size=10, bold=True)
